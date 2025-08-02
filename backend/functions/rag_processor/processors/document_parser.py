"""
Multi-modal Document Processor for RAG

Handles extraction and processing of various document types including PDFs, images,
Word documents, and plain text files for the RAG system.
"""

import logging
import os
import tempfile
from typing import Dict, List, Any, Optional, Tuple, Union
from datetime import datetime
import hashlib
import mimetypes

try:
    import pypdf
    from pdf2image import convert_from_path
    from docx import Document as DocxDocument
    import pytesseract
    from PIL import Image
    import cv2
    import numpy as np
except ImportError as e:
    logging.warning(f"Document processing dependencies not installed: {e}")
    pypdf = None
    convert_from_path = None
    DocxDocument = None
    pytesseract = None
    Image = None
    cv2 = None
    np = None

# For type hints when dependencies are missing
from typing import TYPE_CHECKING
if TYPE_CHECKING or Image is None:
    from typing import Any as Image

from config import Config

logger = logging.getLogger(__name__)

class ExpenseDocumentParser:
    """
    Multi-modal document processor for extracting text and metadata
    """
    
    def __init__(self):
        """Initialize document processor"""
        self.supported_types = {
            'application/pdf': self._process_pdf,
            'image/jpeg': self._process_image,
            'image/jpg': self._process_image,
            'image/png': self._process_image,
            'image/tiff': self._process_image,
            'image/bmp': self._process_image,
            'image/webp': self._process_image,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/msword': self._process_doc,
            'text/plain': self._process_text,
            'text/csv': self._process_text,
            'application/json': self._process_text
        }
        
        # OCR configuration
        self.ocr_config = getattr(Config, 'OCR_CONFIG', '--oem 3 --psm 6')
        
        logger.info(f"Initialized DocumentProcessor with {len(self.supported_types)} supported types")
    
    def process_document(self, 
                        file_path: str,
                        business_id: str,
                        document_id: str = None,
                        metadata: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        Process a document and extract text content with metadata
        
        Args:
            file_path: Path to the document file
            business_id: Business ID for context
            document_id: Optional document ID
            metadata: Additional metadata
            
        Returns:
            Dictionary with extracted content and metadata
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Document not found: {file_path}")
            
            # Generate document ID if not provided
            if not document_id:
                document_id = self._generate_document_id(file_path)
            
            # Detect file type - use mime_type from metadata if available
            mime_type = metadata.get('mime_type') if metadata else None
            if not mime_type:
                mime_type = self._detect_mime_type(file_path)
            
            if mime_type not in self.supported_types:
                raise ValueError(f"Unsupported file type: {mime_type}")
            
            logger.info(f"Processing document: {file_path} (type: {mime_type})")
            
            # Process based on file type
            processor = self.supported_types[mime_type]
            content_data = processor(file_path)
            
            # Combine with metadata
            result = {
                'document_id': document_id,
                'business_id': business_id,
                'file_path': file_path,
                'mime_type': mime_type,
                'file_size': os.path.getsize(file_path),
                'processed_at': datetime.now().isoformat(),
                'content_data': content_data,
                'metadata': metadata or {}
            }
            
            # Add file-level metadata
            result['metadata'].update({
                'original_filename': os.path.basename(file_path),
                'file_extension': os.path.splitext(file_path)[1].lower(),
                'document_type': self._classify_document_type(file_path, content_data)
            })
            
            logger.info(f"Successfully processed document: {document_id}")
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            raise
    
    def _detect_mime_type(self, file_path: str) -> str:
        """Detect MIME type of file"""
        mime_type, _ = mimetypes.guess_type(file_path)
        
        if not mime_type:
            # Fallback detection based on file extension
            ext = os.path.splitext(file_path)[1].lower()
            ext_mapping = {
                '.pdf': 'application/pdf',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.tiff': 'image/tiff',
                '.tif': 'image/tiff',
                '.bmp': 'image/bmp',
                '.webp': 'image/webp',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.doc': 'application/msword',
                '.txt': 'text/plain',
                '.csv': 'text/csv',
                '.json': 'application/json'
            }
            mime_type = ext_mapping.get(ext, 'application/octet-stream')
        
        return mime_type
    
    def _generate_document_id(self, file_path: str) -> str:
        """Generate unique document ID based on file path and content"""
        # Use file path and modification time for uniqueness
        stat = os.stat(file_path)
        content = f"{file_path}_{stat.st_mtime}_{stat.st_size}"
        return hashlib.md5(content.encode()).hexdigest()
    
    def _classify_document_type(self, file_path: str, content_data: Dict[str, Any]) -> str:
        """Classify document type based on filename and content"""
        filename = os.path.basename(file_path).lower()
        
        # Check for common expense document patterns
        if any(word in filename for word in ['receipt', 'invoice', 'bill']):
            return 'expense_document'
        elif any(word in filename for word in ['statement', 'bank']):
            return 'financial_statement'
        elif any(word in filename for word in ['contract', 'agreement']):
            return 'contract'
        elif any(word in filename for word in ['report', 'summary']):
            return 'report'
        
        # Check content for patterns
        text_content = content_data.get('text', '').lower()
        if any(word in text_content for word in ['total:', 'amount:', '$', 'payment', 'transaction']):
            return 'expense_document'
        
        return 'general_document'
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF document"""
        if not pypdf:
            raise ImportError("pypdf not installed. Run: pip install pypdf")
        
        try:
            content_data = {
                'text': '',
                'pages': [],
                'metadata': {},
                'processing_method': 'pdf_extraction'
            }
            
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    content_data['metadata'] = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                        'modification_date': str(pdf_reader.metadata.get('/ModDate', ''))
                    }
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        
                        page_data = {
                            'page_number': page_num + 1,
                            'text': page_text,
                            'char_count': len(page_text)
                        }
                        
                        content_data['pages'].append(page_data)
                        content_data['text'] += f"\n\n--- Page {page_num + 1} ---\n{page_text}"
                        
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        
                        # Try OCR fallback if available
                        if convert_from_path and pytesseract:
                            try:
                                page_image = self._pdf_page_to_image(file_path, page_num)
                                if page_image:
                                    ocr_text = self._ocr_image(page_image)
                                    page_data = {
                                        'page_number': page_num + 1,
                                        'text': ocr_text,
                                        'char_count': len(ocr_text),
                                        'extraction_method': 'ocr_fallback'
                                    }
                                    content_data['pages'].append(page_data)
                                    content_data['text'] += f"\n\n--- Page {page_num + 1} (OCR) ---\n{ocr_text}"
                            except Exception as ocr_error:
                                logger.warning(f"OCR fallback failed for page {page_num + 1}: {str(ocr_error)}")
            
            content_data['total_pages'] = len(content_data['pages'])
            content_data['total_chars'] = len(content_data['text'])
            
            return content_data
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            raise
    
    def _process_image(self, file_path: str) -> Dict[str, Any]:
        """Process image document using OCR"""
        if not pytesseract or not Image:
            raise ImportError("OCR dependencies not installed. Run: pip install pytesseract pillow")
        
        try:
            # Load and preprocess image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Apply image preprocessing for better OCR
            processed_image = self._preprocess_image_for_ocr(image)
            
            # Extract text using OCR
            text = pytesseract.image_to_string(processed_image, config=self.ocr_config)
            
            # Get additional OCR data
            ocr_data = pytesseract.image_to_data(processed_image, output_type=pytesseract.Output.DICT)
            
            content_data = {
                'text': text,
                'processing_method': 'ocr',
                'image_info': {
                    'width': image.width,
                    'height': image.height,
                    'mode': image.mode,
                    'format': image.format
                },
                'ocr_confidence': self._calculate_ocr_confidence(ocr_data),
                'total_chars': len(text)
            }
            
            return content_data
            
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {str(e)}")
            raise
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process Word document"""
        if not DocxDocument:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            full_text = ''
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
                    full_text += para.text + '\n'
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append('\t'.join(row_text))
                tables_text.append('\n'.join(table_text))
                full_text += '\n' + '\n'.join(table_text) + '\n'
            
            content_data = {
                'text': full_text,
                'paragraphs': paragraphs,
                'tables': tables_text,
                'processing_method': 'docx_extraction',
                'total_paragraphs': len(paragraphs),
                'total_tables': len(tables_text),
                'total_chars': len(full_text)
            }
            
            return content_data
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            raise
    
    def _process_doc(self, file_path: str) -> Dict[str, Any]:
        """Process legacy Word document (not implemented - requires additional libraries)"""
        # For .doc files, would need python-docx2txt or similar
        # For now, return empty content with error
        return {
            'text': '',
            'processing_method': 'doc_unsupported',
            'error': 'Legacy .doc format not supported. Please convert to .docx',
            'total_chars': 0
        }
    
    def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            content_data = {
                'text': text,
                'processing_method': 'text_file',
                'total_chars': len(text),
                'total_lines': len(text.splitlines())
            }
            
            return content_data
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                
                content_data = {
                    'text': text,
                    'processing_method': 'text_file_latin1',
                    'total_chars': len(text),
                    'total_lines': len(text.splitlines())
                }
                
                return content_data
                
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {str(e)}")
                raise
    
    def _pdf_page_to_image(self, pdf_path: str, page_num: int) -> Optional[Any]:
        """Convert PDF page to image for OCR"""
        if not convert_from_path:
            return None
        
        try:
            images = convert_from_path(pdf_path, first_page=page_num + 1, last_page=page_num + 1)
            return images[0] if images else None
        except Exception as e:
            logger.warning(f"Error converting PDF page to image: {str(e)}")
            return None
    
    def _preprocess_image_for_ocr(self, image: Any) -> Any:
        """Preprocess image for better OCR results"""
        if not cv2 or not np:
            return image
        
        try:
            # Convert PIL image to OpenCV format
            img_array = np.array(image)
            
            # Convert to grayscale
            if len(img_array.shape) == 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array
            
            # Apply denoising
            denoised = cv2.fastNlMeansDenoising(gray)
            
            # Apply adaptive thresholding
            thresh = cv2.adaptiveThreshold(
                denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
            )
            
            # Convert back to PIL Image
            return Image.fromarray(thresh)
            
        except Exception as e:
            logger.warning(f"Error preprocessing image: {str(e)}")
            return image
    
    def _ocr_image(self, image: Any) -> str:
        """Extract text from image using OCR"""
        if not pytesseract:
            return ""
        
        try:
            return pytesseract.image_to_string(image, config=self.ocr_config)
        except Exception as e:
            logger.warning(f"OCR extraction failed: {str(e)}")
            return ""
    
    def _calculate_ocr_confidence(self, ocr_data: Dict[str, List]) -> float:
        """Calculate average OCR confidence score"""
        try:
            confidences = [int(conf) for conf in ocr_data.get('conf', []) if int(conf) > 0]
            return sum(confidences) / len(confidences) if confidences else 0.0
        except Exception:
            return 0.0
    
    def get_supported_types(self) -> List[str]:
        """Get list of supported MIME types"""
        return list(self.supported_types.keys())
    
    def is_supported(self, file_path: str) -> bool:
        """Check if file type is supported"""
        mime_type = self._detect_mime_type(file_path)
        return mime_type in self.supported_types
    
    def health_check(self) -> Dict[str, Any]:
        """Perform health check on document processor"""
        try:
            status = 'healthy'
            warnings = []
            errors = []
            
            # Check dependencies
            if not pypdf:
                warnings.append("pypdf not available - PDF processing limited")
            
            if not pytesseract:
                warnings.append("pytesseract not available - OCR disabled")
            
            if not DocxDocument:
                warnings.append("python-docx not available - DOCX processing disabled")
            
            if not cv2:
                warnings.append("opencv not available - image preprocessing disabled")
            
            # Test OCR if available
            if pytesseract and Image:
                try:
                    # Create a simple test image
                    test_image = Image.new('RGB', (100, 50), color='white')
                    pytesseract.image_to_string(test_image)
                except Exception as e:
                    errors.append(f"OCR test failed: {str(e)}")
            
            if errors:
                status = 'unhealthy'
            elif warnings:
                status = 'warning'
            
            return {
                'status': status,
                'supported_types': len(self.supported_types),
                'warnings': warnings,
                'errors': errors,
                'ocr_available': pytesseract is not None,
                'pdf_available': pypdf is not None,
                'docx_available': DocxDocument is not None,
                'image_processing_available': cv2 is not None
            }
            
        except Exception as e:
            return {
                'status': 'unhealthy',
                'error': str(e)
            }